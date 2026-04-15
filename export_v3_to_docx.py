"""
Export PROMPT_Xay_dung_Rule_CTKM_30Shine_v3.md → Word (.docx)
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Colors ──
TEAL = RGBColor(0x2E, 0x75, 0xB6)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
GRAY = RGBColor(0x4D, 0x4D, 0x4D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xD6, 0xE4, 0xF0)

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.find(qn('w:shd'))
    if shd is None:
        shd = cell._element.makeelement(qn('w:shd'), {})
        shading.append(shd)
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')

def style_table(table, header_color='2E75B6'):
    """Style a table with header color and borders."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    # Body rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
            if i % 2 == 0:
                set_cell_shading(cell, 'F2F7FB')
    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = tbl.makeelement(qn('w:tblBorders'), {})
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = tbl.makeelement(qn(f'w:{edge}'), {})
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'B0B0B0')

def add_table_from_rows(doc, headers, rows):
    """Add a formatted table to the document."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.autofit = True
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h.strip()
    # Body
    for i, row_data in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i + 1].cells[j]
            val = row_data[j].strip() if j < len(row_data) else ''
            cell.text = val
    style_table(table)
    return table

def parse_md_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = [c.strip() for c in lines[start_idx].strip().strip('|').split('|')]
    # Skip separator line
    data_start = start_idx + 2
    rows = []
    idx = data_start
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break
        row = [c.strip() for c in line.strip('|').split('|')]
        rows.append(row)
        idx += 1
    return headers, rows, idx

def add_code_block(doc, text):
    """Add a code block as a paragraph with monospace font and light background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK
    # light bg via shading on paragraph
    pPr = p._element.get_or_add_pPr()
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = p._element.makeelement(qn('w:shd'), {})
        pPr.append(shd)
    shd.set(qn('w:fill'), 'F5F5F5')
    shd.set(qn('w:val'), 'clear')

def add_run_with_bold(paragraph, text):
    """Add text to paragraph, handling **bold** markers."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def main():
    with open('PROMPT_Xay_dung_Rule_CTKM_30Shine_v3.md', 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')

    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # ── Default font ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(1)

    # ── Heading styles ──
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = TEAL
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(16)
            hs.paragraph_format.space_before = Pt(18)
        elif level == 2:
            hs.font.size = Pt(13)
            hs.paragraph_format.space_before = Pt(14)
        else:
            hs.font.size = Pt(11)
            hs.paragraph_format.space_before = Pt(10)

    # ── Title ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('QUY CHẾ CHƯƠNG TRÌNH KHUYẾN MẠI (CTKM)\n30SHINE')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = TEAL

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('Version 3.0 | Áp dụng từ: Q2/2026')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.italic = True

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run('Đối tượng: SM, GDKV, Giám sát đào tạo ngành hàng (Supervisor)')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    philosophy = doc.add_paragraph()
    philosophy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = philosophy._element.get_or_add_pPr()
    shd = philosophy._element.makeelement(qn('w:shd'), {})
    pPr.append(shd)
    shd.set(qn('w:fill'), 'E8F0FE')
    shd.set(qn('w:val'), 'clear')
    run = philosophy.add_run('Triết lý: CTKM = công cụ kích hoạt hành vi khách hàng, '
                              'tạo liên kết DV ↔ SP ↔ trải nghiệm — không chỉ giảm giá')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = TEAL

    doc.add_paragraph()  # spacer

    # ── Process lines ──
    i = 0
    in_code_block = False
    code_lines = []

    # Skip header lines already processed
    while i < len(lines):
        line = lines[i]
        if line.startswith('## I.'):
            break
        i += 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block toggle
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # Headings
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            text = re.sub(r'\*\*', '', text)
            doc.add_heading(text, level=1)
            i += 1
            continue

        if stripped.startswith('### '):
            text = stripped[4:].strip()
            text = re.sub(r'\*\*', '', text)
            doc.add_heading(text, level=2)
            i += 1
            continue

        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            text = re.sub(r'\*\*', '', text)
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Horizontal rule → skip
        if stripped == '---':
            i += 1
            continue

        # Table
        if stripped.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            # Check if next line is separator
            next_line = lines[i + 1].strip()
            if re.match(r'\|[\s\-:]+\|', next_line):
                headers, rows, end_idx = parse_md_table(lines, i)
                if headers and rows:
                    add_table_from_rows(doc, headers, rows)
                    doc.add_paragraph()  # spacer
                i = end_idx
                continue
            # Single pipe-row without separator = just text
            pass

        # Blockquote
        if stripped.startswith('> '):
            text = stripped[2:].strip()
            text = re.sub(r'\*\*', '', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            # light bg
            pPr = p._element.get_or_add_pPr()
            shd = p._element.makeelement(qn('w:shd'), {})
            pPr.append(shd)
            shd.set(qn('w:fill'), 'F0F4F8')
            shd.set(qn('w:val'), 'clear')
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_run_with_bold(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            text = m.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            add_run_with_bold(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Bold standalone line (like **#1 — "..."**)
        if stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
            text = stripped[2:-2]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = TEAL
            i += 1
            continue

        # Empty line → small spacer
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        if stripped and not stripped.startswith('|'):
            text = stripped
            text = re.sub(r'`([^`]+)`', r'\1', text)  # remove backticks
            p = doc.add_paragraph()
            add_run_with_bold(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        i += 1

    # ── Save ──
    output = 'Quy_che_CTKM_30Shine_v3.docx'
    doc.save(output)
    print(f'Saved: {output}')

if __name__ == '__main__':
    main()
