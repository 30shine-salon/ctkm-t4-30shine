"""
Export 2 CTKM forms (Glanzen + DrForSkin) to Word (.docx)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEAL = RGBColor(0x2E, 0x75, 0xB6)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
GRAY = RGBColor(0x4D, 0x4D, 0x4D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def shading(element, color_hex):
    pPr = element.get_or_add_pPr() if hasattr(element, 'get_or_add_pPr') else element._element.get_or_add_pPr()
    shd = element.makeelement(qn('w:shd'), {}) if hasattr(element, 'makeelement') else element._element.makeelement(qn('w:shd'), {})
    pPr.append(shd)
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')

def set_cell_bg(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = cell._element.makeelement(qn('w:shd'), {})
        tcPr.append(shd)
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')

def style_table(table, header_color='2E75B6'):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = tbl.makeelement(qn('w:tblBorders'), {})
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = tbl.makeelement(qn(f'w:{edge}'), {})
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'B0B0B0')
    for cell in table.rows[0].cells:
        set_cell_bg(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9.5)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_bg(cell, 'F2F7FB')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

def add_form_table(doc, rows_data):
    """Add a 2-column form table: label | content"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.autofit = True
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12)
    for i, (label, content) in enumerate(rows_data):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(label)
        run0.font.bold = True
        run0.font.size = Pt(9.5)
        run0.font.color.rgb = TEAL
        set_cell_bg(c0, 'F0F4F8')
        p1 = c1.paragraphs[0]
        run1 = p1.add_run(content)
        run1.font.size = Pt(9.5)
        run1.font.color.rgb = DARK
    # borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = tbl.makeelement(qn('w:tblBorders'), {})
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = tbl.makeelement(qn(f'w:{edge}'), {})
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'B0B0B0')
    return table

def add_checklist(doc, items):
    for check, text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1)
        symbol = '☑' if check else '☐'
        run = p.add_run(f'  {symbol}  {text}')
        run.font.size = Pt(9.5)
        if check:
            run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        else:
            run.font.color.rgb = GRAY

def add_section_header(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL

def add_note_box(doc, text, bg='FFF8E1'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = DARK
    pPr = p._element.get_or_add_pPr()
    shd = p._element.makeelement(qn('w:shd'), {})
    pPr.append(shd)
    shd.set(qn('w:fill'), bg)
    shd.set(qn('w:val'), 'clear')

def add_comparison_table(doc):
    headers = ['', 'CTKM T3 (cũ)', 'CTKM mới (theo Rule v3)']
    rows = [
        ['Glanzen', 'Nhóm E: "Mua sáp tặng xịt dưỡng"\n→ 2 lượt / toàn hệ thống', 'Nhóm C+F: Trial tạo kiểu trên tóc KH\n→ mua giảm 20%'],
        ['DrForSkin', 'Nhóm E: "Voucher SP 550K"\n→ 1 lượt / toàn hệ thống', 'Nhóm C+E: Mua DrForSkin tặng Apteka\nĐiều kiện: có DV chăm da trong bill'],
        ['Vấn đề cũ', 'KH tự mua SP, không có lý do\nKhông gắn với DV', 'SP gắn với DV → đúng moment\nTrial trực tiếp trên KH'],
        ['Trigger', 'Không có (KH tự biết)', 'Stylist/Skinner chủ động trong lúc làm DV'],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.autofit = True
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    style_table(table)

def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = TEAL
        hs.font.bold = True
        hs.font.size = Pt([0, 18, 14, 11][level])

    # ══════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ĐỀ XUẤT CHƯƠNG TRÌNH KHUYẾN MẠI\n30SHINE — Q2/2026')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = TEAL

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Áp dụng theo Quy chế CTKM v3.0 | 2 chương trình: Sáp Glanzen & Serum DrForSkin 7in1')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.italic = True

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # CT1 — GLANZEN
    # ══════════════════════════════════════════
    doc.add_heading('CT1 — SÁP GLANZEN', level=1)

    add_form_table(doc, [
        ('Tên chương trình', '"Giữ Form — Sáp Glanzen giá đặc biệt khi cắt tạo kiểu"'),
        ('Nhóm CTKM', '☑ C  Cross DV↔SP    ☑ F  Trial / Trải nghiệm'),
        ('Tại sao làm?',
         '• CTKM T3 "Mua sáp tặng xịt dưỡng" = 2 lượt/toàn hệ thống → Nhóm E thuần, thất bại\n'
         '• SC = 65K gói/tháng, base KH cắt rất lớn\n'
         '• Sau cắt stylist vuốt sáp tạo kiểu = moment giới thiệu SP tự nhiên nhất'),
        ('Mục tiêu KD',
         '(Cần xác nhận)\n'
         '☐ Tăng DTT nhóm SP\n'
         '☐ Giải phóng tồn kho Glanzen\n'
         '☐ Giới thiệu SP mới / tăng awareness\n'
         '☐ Tăng cross-sell rate DV cắt → SP tạo kiểu\n'
         '☐ Tạo thói quen chăm tóc tại nhà → KH quay lại cắt đúng kiểu'),
    ])

    doc.add_paragraph()
    add_section_header(doc, 'Cơ chế chi tiết', level=2)

    add_form_table(doc, [
        ('Điều kiện', 'KH sử dụng DV cắt: SC, Cắt xả tạo kiểu, ShineCombo bất kỳ'),
        ('Bước 1 — Trial',
         'Stylist dùng sáp Glanzen tạo kiểu cho KH ngay sau cắt.\n'
         'KH trải nghiệm trực tiếp trên tóc mình — không cần cam kết mua.\n'
         'Lưu ý: Stylist CHỦ ĐỘNG làm luôn, không hỏi "anh có muốn thử không?"'),
        ('Bước 2 — Ưu đãi',
         'Nếu KH thích → mua sáp Glanzen giảm 20%\n'
         'Chỉ áp dụng ngay bill đó (không tặng voucher mua sau)'),
        ('Giới hạn', '1 lần/KH/ngày'),
    ])

    doc.add_paragraph()
    add_note_box(doc,
        '💬 Câu giới thiệu mẫu cho Stylist:\n'
        '"Em vừa vuốt sáp Glanzen tạo kiểu cho anh — anh thấy giữ form không? '
        'Sáp này hôm nay giảm 20% nếu anh muốn dùng ở nhà, '
        'chỉ cần lấy 1 lượng nhỏ là giữ nếp cả ngày."')

    doc.add_paragraph()
    add_section_header(doc, 'Thông tin áp dụng', level=2)

    add_form_table(doc, [
        ('Salon áp dụng', '________________________________    Số salon: ____'),
        ('Thời gian', 'Từ ___/___/2026    Đến ___/___/2026    (tối đa 30 ngày)'),
        ('Người đề xuất', '________________________________    Ngày: ___/___/2026'),
    ])

    doc.add_paragraph()
    add_section_header(doc, 'Kiểm tra nhanh', level=2)

    add_checklist(doc, [
        (True, 'Giảm giá ≤ 40% SP?  →  20% SP = OK'),
        (True, 'Tặng kèm ≤ 100K/bill?  →  Trial tạo kiểu ≈ 0đ = OK'),
        (True, 'Không trùng CTKM HQ đang chạy?'),
        (True, 'Salon đang < 3 CTKM đồng thời?'),
        (False, 'Đã brief stylist/skinner?  →  CẦN BRIEF TRƯỚC KHI CHẠY'),
    ])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('→  5/5 tick + ưu đãi ≤ 100K/bill + ≤ 5 salon  =  TỰ QUYẾT')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL

    doc.add_paragraph()
    add_section_header(doc, 'KPI đo lường', level=2)

    kpi_table = doc.add_table(rows=4, cols=4)
    kpi_table.autofit = True
    kpi_headers = ['KPI', 'Baseline T3', 'Target', 'Cách đo']
    for j, h in enumerate(kpi_headers):
        kpi_table.rows[0].cells[j].text = h
    kpi_data = [
        ['Số lượt mua Glanzen/salon/tuần', '~0', '≥ 5 lượt', 'Hệ thống campaign'],
        ['Tổng lượt mua/tháng/all salon', '2 lượt', '≥ 25 lượt', 'Hệ thống campaign'],
        ['Tỷ lệ convert (KH cắt → mua sáp)', '~0%', '≥ 5%', 'Lượt mua / Tổng bill SC'],
    ]
    for i, row in enumerate(kpi_data):
        for j, val in enumerate(row):
            kpi_table.rows[i+1].cells[j].text = val
    style_table(kpi_table)

    # ══════════════════════════════════════════
    # PAGE BREAK
    # ══════════════════════════════════════════
    doc.add_page_break()

    # ══════════════════════════════════════════
    # CT2 — DRFORSKIN + APTEKA
    # ══════════════════════════════════════════
    doc.add_heading('CT2 — SERUM DRFORSKIN 7IN1 + TẨY DA CHẾT APTEKA', level=1)

    add_form_table(doc, [
        ('Tên chương trình', '"Chăm Da Trọn Bộ — Mua Serum DrForSkin 7in1 tặng Tẩy da chết Apteka"'),
        ('Nhóm CTKM', '☑ C  Cross DV↔SP    ☑ E  Kích hoạt SP'),
        ('Tại sao làm?',
         '• CTKM T3 "Voucher SP 550K" = 1 lượt/toàn hệ thống → thất bại\n'
         '• Nhóm DV chăm da cơ bản: 5,307 gói, SaleRate 4.7%, tăng 3.8% → base KH ổn định\n'
         '• KH vừa làm DV chăm da = da sạch nhất → moment giới thiệu SP dưỡng da tại nhà\n'
         '• Tặng tẩy da chết Apteka → KH có routine mini tại nhà (tẩy + dưỡng)'),
        ('Mục tiêu KD',
         '(Cần xác nhận)\n'
         '☐ Tăng DTT nhóm SP\n'
         '☐ Giới thiệu 2 SP (DrForSkin + Apteka) cho KH\n'
         '☐ Giải phóng tồn kho Apteka (tặng kèm = cách xả thông minh)\n'
         '☐ Tạo routine chăm da tại nhà → KH quay lại DV chăm da thường xuyên\n'
         '☐ Tăng cross-sell rate DV chăm da → SP skincare'),
    ])

    doc.add_paragraph()
    add_section_header(doc, 'Cơ chế chi tiết', level=2)

    add_form_table(doc, [
        ('Điều kiện',
         'KH sử dụng DV chăm da trong bill:\n'
         '• Mặt nạ tẩy da chết sủi bọt\n'
         '• Mặt nạ lạnh Hàn Quốc\n'
         '• Đánh bay mụn cám / lột mụn Full Face\n'
         '• Detox tẩy da chết da đầu\n'
         '• Combo 3 / Combo 5 Chăm da\n'
         '• Shine Spa 199K'),
        ('Ưu đãi',
         'Mua Serum DrForSkin 7in1 → TẶNG Tẩy da chết Apteka\n'
         'Chỉ áp dụng khi bill có DV chăm da (không bán kèm quà khi mua đơn lẻ)'),
        ('Giới hạn',
         '1 combo/KH/ngày\n'
         'Áp dụng: Giá vốn Apteka tặng kèm phải ≤ 80K (theo ngưỡng rule v3)'),
    ])

    doc.add_paragraph()
    add_note_box(doc,
        '💬 Câu giới thiệu mẫu cho Skinner:\n'
        '"Da anh vừa được làm sạch sâu rồi — em thoa lớp serum DrForSkin 7in1 này dưỡng ẩm, '
        'kiềm dầu, chống nắng luôn cho anh. Anh thấy da mịn không? '
        'Hôm nay anh mua 1 tuýp DrForSkin em tặng thêm tẩy da chết Apteka — '
        'tuần tẩy 1 lần rồi thoa serum, da sạch đẹp liên tục."')

    doc.add_paragraph()
    add_section_header(doc, 'Thông tin áp dụng', level=2)

    add_form_table(doc, [
        ('Salon áp dụng', '________________________________    Số salon: ____\n(Ưu tiên salon có SaleRate DV chăm da cao, DV da cơ bản SaleRate 4.7%)'),
        ('Thời gian', 'Từ ___/___/2026    Đến ___/___/2026    (tối đa 30 ngày)'),
        ('Người đề xuất', '________________________________    Ngày: ___/___/2026'),
    ])

    doc.add_paragraph()
    add_section_header(doc, 'Kiểm tra nhanh', level=2)

    add_checklist(doc, [
        (True, 'Giảm giá ≤ 40% SP?  →  Không giảm giá, tặng SP kèm = OK'),
        (True, 'Tặng kèm ≤ 100K/bill?  →  Giá vốn Apteka cần xác nhận ≤ 80K'),
        (True, 'Không trùng CTKM HQ đang chạy?'),
        (True, 'Salon đang < 3 CTKM đồng thời?'),
        (False, 'Đã brief skinner?  →  CẦN BRIEF TRƯỚC KHI CHẠY'),
    ])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('→  5/5 tick + ưu đãi ≤ 100K/bill + ≤ 5 salon  =  TỰ QUYẾT')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL

    doc.add_paragraph()
    add_section_header(doc, 'KPI đo lường', level=2)

    kpi_table2 = doc.add_table(rows=4, cols=4)
    kpi_table2.autofit = True
    for j, h in enumerate(kpi_headers):
        kpi_table2.rows[0].cells[j].text = h
    kpi_data2 = [
        ['Số combo DrForSkin+Apteka / salon / tuần', '~0', '≥ 3 combo', 'Hệ thống campaign'],
        ['Tổng combo/tháng/all salon', '1 lượt', '≥ 15 combo', 'Hệ thống campaign'],
        ['Tỷ lệ convert (KH DV chăm da → mua)', '~0%', '≥ 8%', 'Lượt mua / Bill có DV chăm da'],
    ]
    for i, row in enumerate(kpi_data2):
        for j, val in enumerate(row):
            kpi_table2.rows[i+1].cells[j].text = val
    style_table(kpi_table2)

    # ══════════════════════════════════════════
    # PAGE BREAK — SO SÁNH
    # ══════════════════════════════════════════
    doc.add_page_break()

    doc.add_heading('SO SÁNH CŨ vs MỚI', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Cả 2 CTKM T3 đều thất bại vì dùng Nhóm E (SP thuần, không gắn DV). '
                     'Phiên bản mới chuyển sang Nhóm C (Cross DV↔SP) — gắn SP vào đúng moment KH đang quan tâm.')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.italic = True

    doc.add_paragraph()
    add_comparison_table(doc)

    doc.add_paragraph()

    add_note_box(doc,
        'Điểm chung: KH không cần "quyết định mua" — KH trải nghiệm SP trước '
        '(stylist vuốt sáp / skinner thoa serum), thấy kết quả trên mình, rồi mới được giới thiệu mua. '
        'Đây là khác biệt lớn nhất so với cách cũ.', bg='E8F5E9')

    # ── Save ──
    output = 'De_xuat_CTKM_Glanzen_DrForSkin.docx'
    doc.save(output)
    print(f'Saved: {output}')

if __name__ == '__main__':
    main()
